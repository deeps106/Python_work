# opening a document
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # specifies paragraph justification (e.g.centre, left or right)
from docx.shared import RGBColor  # imports library for coloured text
from docx.shared import Pt  # imports library to change font typeface and size
document = Document()  # opens a blank document based on default template

# adding a heading

heading_1 = document.add_heading("top level heading for document")

heading_2 = document.add_heading("second level heading", level=2)
heading_3 = document.add_heading("3rd level heading", level=3)
 
 
# creating a heading using paragraphs
paragraph_0 = document.add_paragraph()
 
run_0 = paragraph_0.add_run("This is the heading for the document")
run_0.font.bold = True
run_0_format = paragraph_0.paragraph_format
run_0_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

#font size of heading
run_0.font.size = Pt(14)
run_0.font.bold = True
run_0.font.color.rgb = RGBColor(245, 114, 3)  # heading is now in orange




paragraph_1 = document.add_paragraph()  # adding a paragraph

paragraph_format = paragraph_1.paragraph_format  # formats paragraph justification
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # justifies the paragraph to the centre

# using  paragraph to make a heading

# adding bold format to paragraph sentence - need to divide it into runs

run_1 = paragraph_1.add_run("This is a sentence within a new paragraph. ")
run_1.font.size = Pt(20)  # alters font size to 20 for ist pgh

run_2 = paragraph_1.add_run("The second part  of the sentence is in bold. ")
run_2.bold = True  # makes the second sentence  of the paragraph in bold
run_2.font.size = Pt(10)
run_2.font.underline = True  # underlines the sentence
#applying a character style: can be specified when adding a new run
run_3 = paragraph_1.add_run("The third sentence is with emphasis.") # emphasis is italics
run_3.style = "Emphasis"
# align a paragraph to the centre need to import WD_ALIGN_PARAGRAPH (see above)

paragraph_2 = document.add_paragraph()  # adding a paragraph
paragraph_format = paragraph_2.paragraph_format  # formats paragraph justification
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p2_run_1 = paragraph_2.add_run("This second paragraph is in italics with it justified to the right. ")
p2_run_1.style = "Emphasis"

p2_run_2 = paragraph_2.add_run("The right justification is to show that it can be done with paragraph format to the right.")
p2_run_2.style = "Emphasis"

p2_run_2.font.color.rgb = RGBColor(255,0,0) # changes color to red
# saving a document
document.save('test.docx')

# colour format
# import colour format - see above(top)



                                    