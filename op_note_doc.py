import tkinter as tk  # creat a gui for the programme

#creting the ms word document and imports needed
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # specifies paragraph justification (e.g.centre, left or right)
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER  # imports tab library to use in document
from docx.shared import Cm  # uses cm in defining distances
from docx.shared import RGBColor  # imports library for coloured text
from docx.shared import Pt  # imports library to change font typeface and size

document = Document()  # opens a blank document based on default template

# create filename
#filename = input("Give your document a name for saving")

# lists
titles_list = ["Name", "Date of Birth", "Gender", "MRN", "NHS Number", "ASA", "Date", "Theatre", 
                  "First surgeon", "Second surgeon", " Assistant", "Anaesthetist"]

# dictionaries
patient_details_dict = {}

staff_details_dict = {}

def create_dictionaries_from_titles_list():
    for i in range(len(titles_list)):
        if i < 5:
            patient_details_value = input(f"what is the patient's {titles_list[i]}?\n")    
            patient_details_value = patient_details_value.title()
            patient_details_dict.update({titles_list[i]: patient_details_value})
        elif i == 5:
            patient_details_value = input(f"""what is the patient's {titles_list[i]}
                                          1.
                                          2.
                                          3.
                                          4.
                                          5. ?\n""")    
            patient_details_value = patient_details_value.title()
            patient_details_dict.update({titles_list[i]: patient_details_value})
        elif i > 6:
            staff_details_value = input(f"What is the name of the {titles_list[i]}?\n")
            staff_details_value = staff_details_value.title()
            staff_details_dict.update({titles_list[i]: staff_details_value})
            
create_dictionaries_from_titles_list()

# print(patient_details_dict)
# print(staff_details_dict)

# creating the document and populating it with content

# set document font and size
style = document.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(12)
# create the heading operation note in the header of document- level 1
header_section = document.sections[0]  # the first section on document page = header

header = header_section.header  # in first section of document a header is created

# formatting text in header
header_text = header.paragraphs[0]  # first paragraph in header section is the first header/title to which text is added
header_text_format = header_text.paragraph_format
#tab stops
tab_stops = header_text_format.tab_stops
tab_stop = tab_stops.add_tab_stop(Cm(0.25))
# header_text.text = "ROYAL WOLVERHAMPTON NHS TRUST \t\t OPERATION NOTE" # title of document
header_text.add_run("ROYAL WOLVERHAMPTON NHS TRUST\t\t OPERATION NOTE").bold = True


# Add title
title = document.add_heading("OPERATION NOTE")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # justifies title to centre


# add first  paragraph - patient details:
paragraph_0 = document.add_paragraph()
paragraph_format = paragraph_0.paragraph_format
# tab stops
tab_stops = paragraph_format.tab_stops
tab_stop = tab_stops.add_tab_stop(Cm(0.25))

def dictionary_iterate(dict):
    for key, value in dict.items():
        run_0 = paragraph_0.add_run(f"{key}:\t")
        run_0.bold = True
        run_0 = paragraph_0.add_run(f"{value}\t\t")
        run_0.font.style = Pt(10)
dictionary_iterate(patient_details_dict)

# add second paragraph staff details
paragraph_1 = document.add_paragraph()
paragraph_1_format = paragraph_1.paragraph_format
tab_stops = paragraph_1_format.tab_stops
tab_stop = tab_stops.add_tab_stop(Cm(0.5))

for key, value in staff_details_dict.items():
    run_1 = paragraph_1.add_run(f"{key}:\t")
    run_1.bold = True
    run_1 = paragraph_1.add_run(f"{value}\t\t")
    run_1.font.style = Pt(10)

#Anaesthetic
paragraph_2 = document.add_paragraph()
paragraph_2_format = paragraph_2.paragraph_format

while True:
    anaesthetic = input("""Choose from one of the following options for anaesthesia:
                    1. GA
                    2. GA with Spinal
                    3. LA
                    4. Epidural
                    5. Sedation and LA
                    """)
    anaesthetic = int(anaesthetic)
    if  anaesthetic == 1:
        run_2 = paragraph_2.add_run("Anaesthetic:\t")
        run_2.bold = True
        run_2 = paragraph_2.add_run("GA\t\t")
        run_2.font.style = Pt(10)
        break
    elif anaesthetic == 2:
        run_2 = paragraph_2.add_run("Anaesthetic:\t")
        run_2.bold = True
        run_2 = paragraph_2.add_run("GA with Spinal\t\t")
        run_2.font.style = Pt(10)
        break
    elif anaesthetic == 3:
        run_2 = paragraph_2.add_run("Anaesthetic:\t")
        run_2.bold = True
        run_2 = paragraph_2.add_run("LA\t\t")
        run_2.font.style = Pt(10)
        break
    elif anaesthetic == 4:
        run_2 = paragraph_2.add_run("Anaesthetic:\t")
        run_2.bold = True
        run_2 = paragraph_2.add_run("Epidural\t\t")
        run_2.font.style = Pt(10)
        break
    elif anaesthetic == 5:
        run_2 = paragraph_2.add_run("Anaesthetic:\t")
        run_2.bold = True
        run_2 = paragraph_2.add_run("Sedation and LA")
        run_2.font.style = Pt(10)
    
    else:
        print("Shucks! You must've entered text or the wrong number!!. Please try again!")
    
    


document.save("opnote.docx")