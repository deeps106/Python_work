import tkinter as tk  # creat a gui for the programme

from docx import Document
from docx.shared import Inches  # uses inches in defining distances
from docx.shared import Pt  # selects font size
from docx.shared import Length

# input patient information
patient_name = "bob smith"
patient_name = patient_name.title()

gender = "female"
gender = gender.title()

dob = "11/02/74"

asa = "3"

hospital_no = "n8965432"
hospital_no = hospital_no.title()

nhs_no = "111 111 1111"

# input operative team
first_surgeon = "singh-ranger"  # capitalizes first letter
first_surgeon = first_surgeon.title()

second_surgeon = "tanveer"
second_surgeon = second_surgeon.title()

assistant = "datta"
assistant = assistant.title()

anaesthetist = "Dharmeswaran"
anaesthetist = anaesthetist.title()

document = Document()
style = document.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)

sections = document.sections
section = sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(0.5)


document.add_heading("PATIENT DETAILS", level=2)

patient_info = {"NAME":patient_name, 
                     "DATE OF BIRTH":dob, 
                     "GENDER":"Female", 
                     "ASA":asa,
                     "HOSPITAL NUMBER":hospital_no, 
                     "NHS NO":nhs_no}

# paragraph for patient details
p1 = document.add_paragraph()
paragraph_format = p1.paragraph_format
tab_stops = paragraph_format.tab_stops
tab_stop = tab_stops.add_tab_stop(Inches(0.25))
paragraph_format.line_spacing = 1.5

def dictionary_iterate(dict):
    for key, value in dict.items():
        p1.add_run(f"{key}:\t").bold = True
        p1.add_run(f"{value}\n")
dictionary_iterate(patient_info)

op_team = {"FIRST SURGEON":first_surgeon,
           "SECOND SURGEON":second_surgeon,
           "ASSISTANT":assistant,
           "ANAESTHESIOLOGIST":anaesthetist}

document.add_heading("SURGEONS AND ANAESTHETISTS", level=2)

p2 = document.add_paragraph()
paragraph_format = p2.paragraph_format
tab_stops = paragraph_format.tab_stops
tab_stop = tab_stops.add_tab_stop(Inches(0.25))
paragraph_format.line_spacing = 1.5

for key, value in op_team.items():
    p2.add_run(f"{key}:\t").bold = True
    p2.add_run(f"{value}\n")

document.add_heading("ANAESTHESIA", level=2)




document.save("opnote_columns.docx")