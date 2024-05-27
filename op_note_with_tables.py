# opening a document:
from docx import Document
from docx.shared import Pt  # imports library to change font typeface and size


document = Document()
# set document font and size
style = document.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# add a heading

patient_header = document.add_paragraph()
patient_header.add_run("PATIENT DETAILS").bold = True

# add a table

table = document.add_table(rows=2, cols=6)

name_and_dob_gender = table.rows[0]
name_and_dob_gender.cells[0].paragraphs[0].add_run("NAME").bold = True  # styles text to bold
name_and_dob_gender.cells[1].text = "Bob Smith"
name_and_dob_gender.cells[2].paragraphs[0].add_run("DATE OF BIRTH:").bold = True
name_and_dob_gender.cells[3].text = "11/02/1974"
name_and_dob_gender.cells[4].paragraphs[0].add_run("GENDER:").bold = True
name_and_dob_gender.cells[5].text = "Male"


mrn_and_nhs_no_asa= table.rows[1]
mrn_and_nhs_no_asa.cells[0].paragraphs[0].add_run("MRN:").bold = True
mrn_and_nhs_no_asa.cells[1].text = "H456789"
mrn_and_nhs_no_asa.cells[2].paragraphs[0].add_run("NHS No:").bold = True
mrn_and_nhs_no_asa.cells[3].text = "111 111 1111\t"
mrn_and_nhs_no_asa.cells[4].paragraphs[0].add_run("ASA:").bold = True 
mrn_and_nhs_no_asa.cells[5].text = "2"

# staff details
staff_header = document.add_paragraph()
staff_header.add_run("SURGEONS AND ANAESTHETISTS").bold = True
# staff table
staff_table = document.add_table(rows=2, cols=6)

surgeon = staff_table.rows[0]
surgeon.cells[0].text = "First Surgeon:"
surgeon.cells[1].text = "Deepak Singh-Ranger"
surgeon.cells[2].text = "Second Surgeon:"
surgeon.cells[3].text = "Shumaila Tanveer"
surgeon.cells[4].text = "Assistant:"
surgeon.cells[5].text = "Jay Shah"

anaesthetist = staff_table.rows[1]
anaesthetist.cells[0].text = "Anaesthetist:"
anaesthetist.cells[1].text = "Dharmeswaran"

# Operation name
operation_name_header = document.add_paragraph()
operation_name_header.add_run("OPERATION").bold = True

operation_title = document.add_paragraph()
operation_title.add_run(input("What was the operation  performed?"))


# Type of anaesthesia, prep and position table
anaesthetic_type_header = document.add_paragraph()
anaesthetic_type_header.add_run("ANAESTHESIA, PREP AND POSITIONING").bold = True

anaesthesia_table = document.add_table(rows=3, cols=2)

anaesthetic_type = anaesthesia_table.rows[0]
anaesthetic_type.cells[0].text = "Anaesthetic:"
anaesthetic_type.cells[1].text = "GA with Spinal"

prep_type = anaesthesia_table.rows[1]
prep_type.cells[0].text = "Prep:"
prep_type.cells[1].text = "Betadine"

patient_position = anaesthesia_table.rows[2]
patient_position.cells[0].text = "Patient Position:"
patient_position.cells[1].text = "Supine"

# DVT prophylaxis table
dvt_prophylaxis_header = document.add_paragraph()
dvt_prophylaxis_header.add_run("DVT PROPHYLAXIS").bold = True

dvt_prophylaxis_table = document.add_table(rows = 1, cols = 2)

dvt_prophylaxis_info = dvt_prophylaxis_table.rows[0]
dvt_prophylaxis_info.cells[0].text = "DVT Prophylaxis"
dvt_prophylaxis_info.cells[1].text = "TEDS and Flotrons"

# Findings

findings_header = document.add_paragraph()
findings_header.add_run("FINDINGS").bold = True

findings = document.add_paragraph()
findings.add_run(input("What were the operative findings?"))

# Procedure

procedure_header = document.add_paragraph()
procedure_header.add_run("PROCEDURE").bold = True

procedure = document.add_paragraph()
procedure.add_run(input("What were steps of the operation?"))

# closure

closure_header = document.add_paragraph()
closure_header.add_run("CLOSURE").bold = True

closure_info = document.add_paragraph()
skin_closure = closure_info.add_run(input("""Which method did you use to close the skin?
                           1. Skin Staples
                           2. Subcuticular suture
                           3. Interrupted skin closure
                           4. Skin glue
                           5. Staples and skin glue
                           6. Interrupted dermal suture with 3/0 vicryl rapide and skin glue"""))
if skin_closure == 1:
    print("Skin staples")
elif skin_closure == 2:
    suture_material = input("Which suture material did you use?")
    print(f"Subcuticular suture with {suture_material}")
elif skin_closure == 3:
    interrupted_suture_technique = input("Which technique and suture material did you use?")
    print(f"Interrupted skin closure: {interrupted_suture_technique}")
elif skin_closure == 4:
    print("Skin glue")
elif skin_closure == 5:
    print("Staples and skin glue")
elif skin_closure == 6:
    print("Interrupted dermal suture with 3/0 vicryl rapide and skin glue")
else:
    new_skin_closure_method = input("You have not made any of the listed choices. Try again or press 'X' to input your own method")
    new_skin_closure_method = new_skin_closure_method.capitalize()
    if new_skin_closure_method == "X":
        closure_info.add_run(input("What method did you use to close the skin?"))




document.save("op_note_with_tables.docx")